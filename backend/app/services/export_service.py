from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.models import ExportArtifact, Job, Transcript, utcnow
from app.services.audio_service import sha256_file
from app.services.errors import WorkflowError

settings = get_settings()


def safe_filename(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|\x00-\x1f]+", "-", value).strip(" .-")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return (cleaned[:120] or fallback).strip()


def _set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        run._element.get_or_add_rPr().append(rtl)


def _metadata(job: Job, transcript: Transcript) -> dict:
    return {
        "video": {
            "id": job.youtube_video_id,
            "title": job.title,
            "url": job.source_url,
            "channel": job.channel,
            "duration_seconds": job.duration_seconds,
            "playlist_index": job.playlist_index,
        },
        "transcription": {
            "provider": transcript.provider,
            "model": transcript.model,
            "language": transcript.language,
            "text": transcript.full_text,
            "paragraphs": transcript.paragraphs_json,
            "utterances": transcript.utterances_json,
            "words": transcript.words_json,
            "request_id": transcript.request_id,
        },
        "processing": {
            "job_id": job.id,
            "created_at": job.created_at.isoformat() if job.created_at else None,
            "completed_at": job.completed_at.isoformat() if job.completed_at else None,
        },
    }


def generate_docx(job: Job, transcript: Transcript, path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.right_margin = section.left_margin
    title = document.add_heading(job.title or "تفريغ فيديو YouTube", level=1)
    _set_rtl(title)
    metadata_lines = [
        f"القناة: {job.channel or '-'}",
        f"رابط الفيديو: {job.source_url}",
        f"اللغة: {transcript.language}",
        f"النموذج: {transcript.model}",
    ]
    for line in metadata_lines:
        paragraph = document.add_paragraph(line)
        _set_rtl(paragraph)
    document.add_section(WD_SECTION.CONTINUOUS)
    for block in [part.strip() for part in transcript.full_text.split("\n") if part.strip()]:
        paragraph = document.add_paragraph(block)
        _set_rtl(paragraph)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def generate_txt(job: Job, transcript: Transcript, path: Path) -> None:
    content = (
        f"{job.title or 'تفريغ فيديو YouTube'}\n"
        f"القناة: {job.channel or '-'}\n"
        f"الرابط: {job.source_url}\n\n"
        f"{transcript.full_text}\n"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def generate_json(job: Job, transcript: Transcript, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(_metadata(job, transcript), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def ensure_exports(db: Session, job: Job) -> list[ExportArtifact]:
    transcript = job.transcript
    if transcript is None:
        raise WorkflowError("EXPORT_NO_TRANSCRIPT", "لا يوجد تفريغ لإنشاء الملفات")
    folder = settings.exports_root / job.id
    base = safe_filename(job.title or job.youtube_video_id or job.id, job.id)
    generators = {
        "docx": (folder / f"{base}.docx", generate_docx),
        "txt": (folder / f"{base}.txt", generate_txt),
        "json": (folder / f"{base}.json", generate_json),
    }
    artifacts: list[ExportArtifact] = []
    by_format = {artifact.format: artifact for artifact in job.exports}
    for format_name, (path, generator) in generators.items():
        artifact = by_format.get(format_name)
        if not path.exists() or path.stat().st_size == 0:
            try:
                generator(job, transcript, path)
            except Exception as exc:
                raise WorkflowError(
                    "EXPORT_FAILED",
                    f"فشل إنشاء ملف {format_name.upper()}",
                    str(exc),
                    True,
                ) from exc
        if artifact is None:
            artifact = ExportArtifact(job_id=job.id, format=format_name, file_path=str(path))
            db.add(artifact)
        artifact.file_path = str(path)
        artifact.file_size = path.stat().st_size
        artifact.checksum = sha256_file(path)
        artifact.created_at = utcnow()
        artifacts.append(artifact)
    db.commit()
    return artifacts
